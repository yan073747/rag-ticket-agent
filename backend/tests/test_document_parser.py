import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import ArrayObject, DecodedStreamObject, DictionaryObject, NameObject

from app.services.document_parser import DocumentParseError, parse_document_text


def write_pdf_with_text(path: Path, text: str) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=200)

    font_resource = DictionaryObject(
        {
            NameObject("/F1"): DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
        }
    )
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): font_resource})

    stream = DecodedStreamObject()
    safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(f"BT /F1 12 Tf 50 150 Td ({safe_text}) Tj ET".encode("utf-8"))
    stream_reference = writer._add_object(stream)
    page[NameObject("/Contents")] = stream_reference

    writer.write(path)


class DocumentParserTest(unittest.TestCase):
    def test_parse_pdf_extracts_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "policy.pdf"
            write_pdf_with_text(pdf_path, "Refunds are handled within 7 business days.")

            text = parse_document_text(pdf_path)

        self.assertIn("Refunds are handled", text)

    def test_parse_docx_extracts_paragraphs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "policy.docx"
            document = Document()
            document.add_paragraph("Annual leave policy")
            document.add_paragraph("Employees receive 10 days of annual leave.")
            document.save(docx_path)

            text = parse_document_text(docx_path)

        self.assertIn("Annual leave policy", text)
        self.assertIn("10 days", text)

    def test_empty_pdf_raises_parse_error(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "empty.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=200)
            writer.write(pdf_path)

            with self.assertRaises(DocumentParseError):
                parse_document_text(pdf_path)


if __name__ == "__main__":
    unittest.main()
